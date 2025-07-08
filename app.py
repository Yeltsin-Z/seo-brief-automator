from flask import Flask, render_template, request, jsonify, send_file
import os
import sys
import logging
from datetime import datetime
import json
import threading
import time
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import markdown2

# Add the project root to the Python path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from src.serp_scraper import SERPScraper
from src.ugc_researcher import UGCResearcher
from src.brief_generator import BriefGenerator
from config.config import Config

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Debug environment variables
logger.info("=== Environment Debug Info ===")
logger.info(f"OPENAI_API_KEY set: {'Yes' if os.getenv('OPENAI_API_KEY') else 'No'}")
if os.getenv('OPENAI_API_KEY'):
    logger.info(f"OPENAI_API_KEY starts with: {os.getenv('OPENAI_API_KEY')[:10]}...")
logger.info(f"Config.OPENAI_API_KEY: {'Set' if Config.OPENAI_API_KEY else 'Not set'}")
logger.info("=== End Debug Info ===")

# Global variable to store current job status
current_job = {
    'status': 'idle',
    'progress': 0,
    'message': '',
    'result': None,
    'error': None,
    'step': 'idle',  # 'serp', 'ugc', 'serp_analysis', 'complete'
    'serp_results': [],
    'ugc_brief': None,
    'serp_brief': None,
    'final_brief': None
}

# Global API request counter
api_request_count = 0
MAX_API_REQUESTS = 200

def increment_api_request_count():
    """Increment the API request counter and check if limit is reached"""
    global api_request_count
    api_request_count += 1
    logger.info(f"API request count: {api_request_count}/{MAX_API_REQUESTS}")
    
    if api_request_count >= MAX_API_REQUESTS:
        logger.warning(f"API request limit reached: {api_request_count}/{MAX_API_REQUESTS}")
        return True  # Limit reached
    return False  # Still under limit

def reset_api_request_count():
    """Reset the API request counter"""
    global api_request_count
    api_request_count = 0
    logger.info("API request counter reset")

def get_api_request_status():
    """Get current API request status"""
    global api_request_count
    return {
        'current_count': api_request_count,
        'max_requests': MAX_API_REQUESTS,
        'limit_reached': api_request_count >= MAX_API_REQUESTS
    }

def run_step_1_serp_scraping(focus_keyword, topic_theme, buyer_persona, content_id):
    """Step 1: Collect top 10 SERP results using SerpAPI"""
    global current_job
    try:
        current_job['status'] = 'running'
        current_job['step'] = 'serp'
        current_job['progress'] = 0
        current_job['message'] = 'Step 1: Collecting top 10 SERP results...'
        current_job['error'] = None
        current_job['focus_keyword'] = focus_keyword
        current_job['topic_theme'] = topic_theme
        current_job['buyer_persona'] = buyer_persona
        current_job['content_id'] = content_id
        # Use SERP scraper to get top 10 results
        scraper = SERPScraper()
        serp_results = scraper.search_google(focus_keyword)
        current_job['serp_results'] = serp_results
        current_job['progress'] = 25
        current_job['message'] = f'Step 1 Complete: Found {len(serp_results)} SERP results. Ready for UGC research.'
        current_job['step'] = 'serp_complete'
        logger.info(f"Step 1 completed: Found {len(serp_results)} SERP results")
    except Exception as e:
        logger.error(f"Error in Step 1 (SERP scraping): {e}")
        current_job['status'] = 'error'
        current_job['error'] = str(e)
        current_job['message'] = f'Step 1 Error: {str(e)}'

def run_step_2_ugc_research(focus_keyword, topic_theme, buyer_persona, custom_prompt):
    """Step 2: Generate UGC brief using basic research"""
    global current_job
    
    try:
        current_job['status'] = 'running'
        current_job['step'] = 'ugc'
        current_job['progress'] = 25
        current_job['message'] = 'Step 2: Conducting comprehensive UGC research...'
        current_job['error'] = None
        
        # Initialize UGC researcher
        ugc_researcher = UGCResearcher(api_increment_callback=increment_api_request_count)
        
        # Perform basic UGC research (faster and simpler)
        ugc_data = ugc_researcher.research_ugc_basic(
            focus_keyword=focus_keyword,
            topic_theme=topic_theme,
            buyer_persona=buyer_persona,
            custom_prompt=custom_prompt
        )
        
        current_job['ugc_brief'] = ugc_data
        current_job['progress'] = 50
        current_job['message'] = 'Step 2 Complete: Enhanced UGC research completed. Ready for SERP analysis.'
        current_job['step'] = 'ugc_complete'
        
        logger.info("Step 2 completed: UGC brief generated")
        
    except Exception as e:
        logger.error(f"Error in Step 2 (UGC research): {e}")
        current_job['status'] = 'error'
        current_job['error'] = str(e)
        current_job['message'] = f'Step 2 Error: {str(e)}'

def run_step_3_serp_analysis(focus_keyword, custom_prompt):
    """Step 3: Generate SERP brief using Prompt #2"""
    global current_job
    
    try:
        current_job['status'] = 'running'
        current_job['step'] = 'serp_analysis'
        current_job['progress'] = 50
        current_job['message'] = 'Step 3: Generating SERP analysis brief...'
        current_job['error'] = None
        
        # Initialize UGC researcher for SERP analysis
        ugc_researcher = UGCResearcher(api_increment_callback=increment_api_request_count)
        
        # Use the specific Prompt #2 for SERP analysis
        serp_prompt = """You are an expert content strategist.
I will give you a list of blog article URLs that rank on the first page of Google for a target keyword.
For each URL:
- Start with the blog title and URL in bold
- Then provide a structured analysis in bullet format under the following sections

Please follow this structure for each blog:
Blog Title: [Insert blog post title]
URL: [Insert blog post URL]
1. Summary of the Article
   - Brief overview of what the article covers.
2. Search Intent Covered
   - What is the primary search intent (Informational, Navigational, Transactional, Commercial)?
   - How well does it fulfill that intent?
3. Subtopics Covered
   - Bullet list of all key sections or talking points.
4. Depth of Coverage
   - Is the content shallow, moderate, or deep?
   - What's covered in detail vs only skimmed?
5. What's Missing
   - Gaps in logic, unexplored ideas, missing data/examples/frameworks.
6. Tone & Point of View
   - Describe the tone (e.g., expert, conversational, tactical).
   - Is there a distinct point of view or is it generic?
7. Structure
   - Type of article (e.g., guide, listicle, editorial).
   - Include the full H1, and a list of H2s and H3s (if present).
8. SEO Signals
   - Is the primary keyword used in title, H1, intro, and headers?
   - Mention 7-10 semantically related keywords used (if any).
   - Describe internal and external linking strategy.
9. Use of Visuals
   - What types of visuals are used (if any)?
   - Are they helpful, decorative, or generic?
10. Unique Hooks or Contrarian Angles
   - Any standout story, metaphor, framework, or angle?
11. Opportunities to Differentiate
   - What can be improved — stronger POV, richer examples, deeper insights, visual storytelling, etc.?
12. Word Count
   - Estimated total word count of the article."""

        # Use custom prompt if provided
        if custom_prompt:
            serp_prompt = custom_prompt
        
        # Perform SERP analysis
        serp_analysis = ugc_researcher.perform_serp_analysis(
            serp_results=current_job['serp_results'],
            focus_keyword=focus_keyword,
            custom_prompt=serp_prompt
        )
        
        current_job['serp_brief'] = serp_analysis
        current_job['progress'] = 75
        current_job['message'] = 'Analysis Complete'
        current_job['step'] = 'serp_analysis_complete'
        
        logger.info("Step 3 completed: SERP analysis brief generated")
        
    except Exception as e:
        logger.error(f"Error in Step 3 (SERP analysis): {e}")
        current_job['status'] = 'error'
        current_job['error'] = str(e)
        current_job['message'] = f'Step 3 Error: {str(e)}'

def run_step_4_combine_results(focus_keyword, topic_theme, buyer_persona, content_id, custom_prompt):
    """Step 4: Generate final brief using ChatGPT with custom prompt"""
    global current_job
    
    try:
        current_job['status'] = 'running'
        current_job['step'] = 'combine'
        current_job['progress'] = 75
        current_job['message'] = 'Step 4: Generating final brief with ChatGPT...'
        current_job['error'] = None
        
        # Check API limit
        global api_request_count
        if api_request_count >= 200:
            raise Exception("API request limit reached (200 requests)")
        
        # Get UGC and SERP data
        ugc_brief = current_job.get('ugc_brief', {})
        serp_brief = current_job.get('serp_brief', {})
        serp_results = current_job.get('serp_results', [])
        
        # Prepare the context for ChatGPT
        ugc_content = ugc_brief.get('raw_response', 'No UGC research available')
        serp_content = serp_brief.get('raw_response', 'No SERP analysis available')
        
        # Create the full prompt with context
        full_prompt = f"{custom_prompt}\n\nUGC Research Document:\n{ugc_content}\n\nSERP Analysis Document:\n{serp_content}"
        
        # Call ChatGPT
        increment_api_request_count()
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert SEO content writer specializing in creating comprehensive content briefs for finance and business topics."},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        # Extract the response
        final_brief_content = response.choices[0].message.content
        
        # Convert markdown to HTML for display
        html_output = markdown2.markdown(final_brief_content, extras=["fenced-code-blocks", "tables", "strike", "cuddled-lists", "metadata", "numbering", "task_list"])
        html_output = html_output.replace('<ul>', '<ul class="custom-list">').replace('<ol>', '<ol class="custom-list-ol">').replace('<li>', '<li class="custom-list-item">')
        
        # Create the final brief result
        final_brief = {
            'focus_keyword': focus_keyword,
            'topic_theme': topic_theme,
            'buyer_persona': buyer_persona,
            'content_id': content_id,
            'generated_at': datetime.now().isoformat(),
            'raw_response': final_brief_content,
            'raw_markdown': final_brief_content,
            'html_output': html_output,
            'status': 'completed',
            'message': 'Final brief generated successfully'
        }
        
        # Save to file
        from src.brief_generator import BriefGenerator
        generator = BriefGenerator(api_increment_callback=increment_api_request_count)
        
        # Create a combined result structure for saving
        combined_result = {
            'focus_keyword': focus_keyword,
            'topic_theme': topic_theme,
            'buyer_persona': buyer_persona,
            'content_id': content_id,
            'generated_at': datetime.now().isoformat(),
            'status': 'completed',
            'serp_analysis': {
                'articles_found': len(serp_results),
                'articles': serp_results
            },
            'ugc_research': ugc_brief,
            'final_brief': final_brief,
            'brief_content': final_brief
        }
        
        filename = generator.save_brief(combined_result)
        final_brief['filename'] = filename
        
        # Update current job
        current_job['final_brief'] = final_brief
        current_job['result'] = combined_result
        current_job['progress'] = 100
        current_job['message'] = 'Complete: Final brief generated successfully!'
        current_job['status'] = 'completed'
        current_job['step'] = 'complete'
        
        logger.info("Step 4 completed: Final brief generated with ChatGPT")
        
    except Exception as e:
        logger.error(f"Error in Step 4 (generate final brief): {e}")
        current_job['status'] = 'error'
        current_job['error'] = str(e)
        current_job['message'] = f'Step 4 Error: {str(e)}'

@app.route('/')
def index():
    """Main page with the refactored brief generator interface"""
    return render_template('index.html')

@app.route('/start-step1', methods=['POST'])
def start_step1():
    """Start Step 1: SERP Collection"""
    global current_job
    
    try:
        data = request.get_json()
        focus_keyword = data.get('focus_keyword', '').strip()
        topic_theme = data.get('topic_theme', '').strip()
        buyer_persona = data.get('buyer_persona', '').strip()
        
        if not focus_keyword or not topic_theme or not buyer_persona:
            return jsonify({'error': 'Focus keyword, topic theme, and buyer persona are required'}), 400
        
        # Generate a content ID automatically
        content_id = f"{focus_keyword[:3].upper()}{len(focus_keyword)}"
        
        # Reset job status and API counter
        current_job = {
            'status': 'idle',
            'progress': 0,
            'message': '',
            'result': None,
            'error': None,
            'step': 'idle',
            'serp_results': [],
            'ugc_brief': None,
            'serp_brief': None,
            'final_brief': None,
            'focus_keyword': focus_keyword,
            'topic_theme': topic_theme,
            'buyer_persona': buyer_persona,
            'content_id': content_id
        }
        reset_api_request_count()
        
        # Start Step 1 in background thread
        thread = threading.Thread(
            target=run_step_1_serp_scraping,
            args=(focus_keyword, topic_theme, buyer_persona, content_id)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'status': 'started',
            'message': 'Step 1: SERP scraping started successfully'
        })
        
    except Exception as e:
        logger.error(f"Error starting Step 1: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/start-step2', methods=['POST'])
def start_step2():
    """Start Step 2: UGC Research"""
    global current_job
    
    try:
        data = request.get_json()
        custom_prompt = data.get('custom_prompt', '').strip()
        
        if current_job['step'] != 'serp_complete':
            return jsonify({'error': 'Step 1 must be completed first'}), 400
        
        # Start Step 2 in background thread
        thread = threading.Thread(
            target=run_step_2_ugc_research,
            args=(current_job['focus_keyword'], current_job['topic_theme'], 
                  current_job['buyer_persona'], custom_prompt)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'status': 'started',
            'message': 'Step 2: UGC research started successfully'
        })
        
    except Exception as e:
        logger.error(f"Error starting Step 2: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/start-step3', methods=['POST'])
def start_step3():
    """Start Step 3: SERP Analysis"""
    global current_job
    
    try:
        data = request.get_json()
        custom_prompt = data.get('custom_prompt', '').strip()
        
        if current_job['step'] != 'ugc_complete':
            return jsonify({'error': 'Step 2 must be completed first'}), 400
        
        # Start Step 3 in background thread
        thread = threading.Thread(
            target=run_step_3_serp_analysis,
            args=(current_job['focus_keyword'], custom_prompt)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'status': 'started',
            'message': 'Step 3: SERP analysis started successfully'
        })
        
    except Exception as e:
        logger.error(f"Error starting Step 3: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/start-step4', methods=['POST'])
def start_step4():
    """Start Step 4: Combine Results"""
    global current_job
    
    try:
        if current_job['step'] != 'serp_analysis_complete':
            return jsonify({'error': 'Step 3 must be completed first'}), 400
        
        # Get custom prompt from request
        data = request.get_json()
        custom_prompt = data.get('custom_prompt', '') if data else ''
        
        # Get content_id from current_job or generate one
        content_id = current_job.get('content_id')
        if not content_id:
            focus_keyword = current_job.get('focus_keyword', '')
            content_id = f"{focus_keyword[:3].upper()}{len(focus_keyword)}"
        
        # Start Step 4 in background thread
        thread = threading.Thread(
            target=run_step_4_combine_results,
            args=(current_job['focus_keyword'], current_job['topic_theme'], 
                  current_job['buyer_persona'], content_id, custom_prompt)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'status': 'started',
            'message': 'Step 4: Generating final brief started successfully'
        })
        
    except Exception as e:
        logger.error(f"Error starting Step 4: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/status')
def get_status():
    """Get current job status"""
    global current_job
    status_data = current_job.copy()
    status_data['api_status'] = get_api_request_status()
    return jsonify(status_data)

@app.route('/download/<filename>')
def download_file(filename):
    """Download generated brief file"""
    try:
        file_path = os.path.join('output', filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-step2-pdf', methods=['POST'])
def download_step2_pdf():
    """Generate and download Step 2 UGC Brief PDF"""
    try:
        from src.pdf_generator import PDFGenerator
        
        if current_job['step'] != 'ugc_complete':
            return jsonify({'error': 'Step 2 must be completed first'}), 400
        
        pdf_generator = PDFGenerator()
        filename = pdf_generator.generate_step2_pdf(
            current_job['ugc_brief'],
            current_job['focus_keyword'],
            current_job['topic_theme'],
            current_job['buyer_persona']
        )
        
        if filename:
            return jsonify({
                'status': 'success',
                'filename': filename,
                'message': 'Step 2 PDF generated successfully'
            })
        else:
            return jsonify({'error': 'Failed to generate PDF'}), 500
            
    except Exception as e:
        logger.error(f"Error generating Step 2 PDF: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-step3-pdf', methods=['POST'])
def download_step3_pdf():
    """Generate and download Step 3 SERP Analysis PDF"""
    try:
        from src.pdf_generator import PDFGenerator
        
        if current_job['step'] != 'serp_analysis_complete':
            return jsonify({'error': 'Step 3 must be completed first'}), 400
        
        pdf_generator = PDFGenerator()
        filename = pdf_generator.generate_step3_pdf(
            current_job['serp_brief'],
            current_job['serp_results'],
            current_job['focus_keyword']
        )
        
        if filename:
            return jsonify({
                'status': 'success',
                'filename': filename,
                'message': 'Step 3 PDF generated successfully'
            })
        else:
            return jsonify({'error': 'Failed to generate PDF'}), 500
            
    except Exception as e:
        logger.error(f"Error generating Step 3 PDF: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-final-pdf', methods=['POST'])
def download_final_pdf():
    """Generate and download Final Comprehensive PDF"""
    try:
        from src.pdf_generator import PDFGenerator
        
        if current_job['step'] != 'complete':
            return jsonify({'error': 'All steps must be completed first'}), 400
        
        pdf_generator = PDFGenerator()
        filename = pdf_generator.generate_final_pdf(current_job['result'])
        
        if filename:
            return jsonify({
                'status': 'success',
                'filename': filename,
                'message': 'Final PDF generated successfully'
            })
        else:
            return jsonify({'error': 'Failed to generate PDF'}), 500
            
    except Exception as e:
        logger.error(f"Error generating final PDF: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-step3-docx', methods=['POST'])
def download_step3_docx():
    """Generate and download Step 3 SERP Analysis DOCX"""
    try:
        if current_job['step'] != 'serp_analysis_complete':
            return jsonify({'error': 'Step 3 must be completed first'}), 400
        serp_brief = current_job.get('serp_brief', {})
        content = serp_brief.get('raw_markdown', '')
        if not content:
            return jsonify({'error': 'No SERP analysis content available'}), 400
        focus_keyword = serp_brief.get('focus_keyword', 'N/A')
        output_dir = 'output'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"SERP_Analysis_{focus_keyword.replace(' ', '_')}_{timestamp}.docx"
        file_path = os.path.join(output_dir, filename)
        doc = Document()
        title = doc.add_heading('SERP Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Focus Keyword: {focus_keyword}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        for line in content.split('\n'):
            line = line.strip()
            if line:
                if line.startswith('#'):
                    level = line.count('#')
                    doc.add_heading(line.lstrip('#').strip(), level=min(level, 4))
                else:
                    doc.add_paragraph(line)
        doc.save(file_path)
        return jsonify({'status': 'success', 'filename': filename, 'message': 'SERP DOCX generated successfully'})
    except Exception as e:
        logger.error(f"Error generating SERP DOCX: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download-final-docx', methods=['POST'])
def download_final_docx():
    """Generate and download Final Comprehensive DOCX"""
    try:
        if current_job['step'] != 'complete':
            return jsonify({'error': 'All steps must be completed first'}), 400
        final_brief = current_job.get('final_brief', {})
        content = final_brief.get('raw_markdown', final_brief.get('content', ''))
        if not content:
            return jsonify({'error': 'No final brief content available'}), 400
        focus_keyword = final_brief.get('focus_keyword', 'N/A')
        topic_theme = final_brief.get('topic_theme', 'N/A')
        buyer_persona = final_brief.get('buyer_persona', 'N/A')
        content_id = final_brief.get('content_id', 'N/A')
        output_dir = 'output'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"SEO_Brief_{focus_keyword.replace(' ', '_')}_{timestamp}.docx"
        file_path = os.path.join(output_dir, filename)
        doc = Document()
        title = doc.add_heading('SEO Content Brief', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Focus Keyword: {focus_keyword}")
        doc.add_paragraph(f"Topic Theme: {topic_theme}")
        doc.add_paragraph(f"Buyer Persona: {buyer_persona}")
        doc.add_paragraph(f"Content ID: {content_id}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        for line in content.split('\n'):
            line = line.strip()
            if line:
                if line.startswith('#'):
                    level = line.count('#')
                    doc.add_heading(line.lstrip('#').strip(), level=min(level, 4))
                else:
                    doc.add_paragraph(line)
        doc.save(file_path)
        return jsonify({'status': 'success', 'filename': filename, 'message': 'Final DOCX generated successfully'})
    except Exception as e:
        logger.error(f"Error generating final DOCX: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/ugc-html')
def get_ugc_html():
    """Get UGC research HTML output"""
    global current_job
    
    try:
        if current_job['step'] != 'ugc_complete':
            return jsonify({'error': 'Step 2 must be completed first'}), 400
        
        ugc_brief = current_job.get('ugc_brief', {})
        html_output = ugc_brief.get('html_output', '')
        
        if html_output:
            return html_output, 200, {'Content-Type': 'text/html'}
        else:
            return jsonify({'error': 'No HTML output available'}), 404
            
    except Exception as e:
        logger.error(f"Error getting UGC HTML: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/serp-html')
def get_serp_html():
    """Get SERP analysis HTML output"""
    global current_job
    try:
        if current_job['step'] != 'serp_analysis_complete':
            return jsonify({'error': 'Step 3 must be completed first'}), 400
        serp_brief = current_job.get('serp_brief', {})
        html_output = serp_brief.get('html_output', '')
        if html_output:
            return html_output, 200, {'Content-Type': 'text/html'}
        else:
            return jsonify({'error': 'No HTML output available'}), 404
    except Exception as e:
        logger.error(f"Error getting SERP HTML: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/final-html')
def get_final_html():
    """Get final brief HTML output"""
    global current_job
    try:
        if current_job['step'] != 'complete':
            return jsonify({'error': 'All steps must be completed first'}), 400
        final_brief = current_job.get('final_brief', {})
        html_output = final_brief.get('html_output', '')
        if html_output:
            return html_output, 200, {'Content-Type': 'text/html'}
        else:
            return jsonify({'error': 'No HTML output available'}), 404
    except Exception as e:
        logger.error(f"Error getting final HTML: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/examples')
def get_examples():
    """Get example keywords"""
    examples = [
        {
            'keyword': 'financial planning',
            'description': 'Comprehensive financial planning strategies',
            'category': 'Finance'
        },
        {
            'keyword': 'budget management',
            'description': 'Effective budget management techniques',
            'category': 'Finance'
        },
        {
            'keyword': 'investment strategies',
            'description': 'Investment strategies for different goals',
            'category': 'Finance'
        }
    ]
    return jsonify(examples)

@app.route('/api/recent-briefs')
def get_recent_briefs():
    """Get recent briefs from output directory"""
    try:
        output_dir = 'output'
        if not os.path.exists(output_dir):
            return jsonify([])
        
        files = []
        for filename in os.listdir(output_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(output_dir, filename)
                stat = os.stat(file_path)
                files.append({
                    'filename': filename,
                    'created_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    'size': stat.st_size
                })
        
        # Sort by creation time, newest first
        files.sort(key=lambda x: x['created_at'], reverse=True)
        return jsonify(files[:10])  # Return last 10 files
        
    except Exception as e:
        logger.error(f"Error getting recent briefs: {e}")
        return jsonify([])

@app.route('/download-step2-docx', methods=['POST'])
def download_step2_docx():
    """Generate and download Step 2 UGC Research DOCX"""
    try:
        if current_job['step'] != 'ugc_complete':
            return jsonify({'error': 'Step 2 must be completed first'}), 400
        
        ugc_brief = current_job.get('ugc_brief', {})
        content = ugc_brief.get('raw_markdown', '')
        if not content:
            return jsonify({'error': 'No UGC research content available'}), 400
        
        focus_keyword = ugc_brief.get('focus_keyword', 'N/A')
        topic_theme = ugc_brief.get('topic_theme', 'N/A')
        buyer_persona = ugc_brief.get('buyer_persona', 'N/A')
        
        output_dir = 'output'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"UGC_Research_{focus_keyword.replace(' ', '_')}_{timestamp}.docx"
        file_path = os.path.join(output_dir, filename)
        
        doc = Document()
        title = doc.add_heading('UGC Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Focus Keyword: {focus_keyword}")
        doc.add_paragraph(f"Topic Theme: {topic_theme}")
        doc.add_paragraph(f"Buyer Persona: {buyer_persona}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        # Add content (markdown to docx, simple line by line)
        for line in content.split('\n'):
            line = line.strip()
            if line:
                if line.startswith('#'):
                    level = line.count('#')
                    doc.add_heading(line.lstrip('#').strip(), level=min(level, 4))
                else:
                    doc.add_paragraph(line)
        doc.save(file_path)
        return jsonify({'status': 'success', 'filename': filename, 'message': 'UGC DOCX generated successfully'})
    except Exception as e:
        logger.error(f"Error generating UGC DOCX: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True) 